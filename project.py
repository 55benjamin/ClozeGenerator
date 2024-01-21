# standard python modules
import re
import time
import random
from os.path import splitext

# installable modules
import spacy
import curses
import docx2txt
from pdfminer.high_level import extract_text
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH



def main(stdscr):
    # wait for the user to select the input mode
    mode = get_mode(stdscr)

    while True: 
        try: 
            source = get_source(stdscr,mode)
            content = get_content(source, mode)
            limit = get_limit(stdscr)
            dest = get_dest(stdscr)
            title = get_title(stdscr)
            break

        except FileNotFoundError:
            error_msg = 'File not found. Please try entering a different name and ensure the extension is correct.'

            show_error(stdscr,error_msg)

            stdscr.refresh()
            time.sleep(2)
            
            continue
    
    # show progress message 
    stdscr.clear()
    stdscr.addstr(0,1, "Generating cloze passage...")
    stdscr.refresh()
    
    # write to the question and answer files 
    if dest == '.docx':
        generate_to_docx(content, title, limit)

    elif dest == '.txt':
        generate_to_txt(content, title, limit)

    stdscr.clear()
    success_msg = f"Cloze passage generated to questions{dest} and answers{dest}!\n Press ENTER to exit."
    show_success(stdscr, success_msg)
    
    # wait for user to press enter to exit the program
    stdscr.getch()


"""
Disclosure: The following functions prefixed with 'get' rely on python's 'curses' module. 
In arriving at the decision to use this module and in understanding some of the functions and attributes in this module, 
I probed ChatGPT based on my objectives and requirements. I did not copy code from it verbatim and only referred to it, in tandem with 
the official 'curses' documentation (https://docs.python.org/3/library/curses.html#), to better understand an error or to narrow down on 
its documentation. The prompts that I used are viewable below:

Clickable Python Terminal: Curses: https://chat.openai.com/share/eaa46c49-4c74-423f-9bbc-5d0d95eb6f17
Decode Curses Input: https://chat.openai.com/share/68af53d0-65f5-4065-baf8-bc36ff7c8e89
"""


# lets user click to select input mode (.pdf, .docx or .txt)
def get_mode(stdscr):

    # clear screen and hide the cursor
    stdscr.clear()
    curses.curs_set(0)

    # add title text 
    stdscr.addstr(0, 1, "Cloze Generator. Updated Jan 2023")
    stdscr.addstr(1, 1, "https://github.com/55benjamin/ClozeGenerator")
    stdscr.refresh()

    # record all mouse movements and get the mouse position
    curses.mousemask(curses.ALL_MOUSE_EVENTS | curses.REPORT_MOUSE_POSITION)

    # option text to be displayed
    pdf_opt = "PDF (.pdf)"
    word_opt = "Word (.docx)"
    text_opt = "Text (.txt)"
    

    # display options on the screen
    stdscr.addstr(3,1, "Step 1 of 5: To start, please select and click on an input method:")
    stdscr.addstr(5,1, pdf_opt)
    stdscr.addstr(5,20, word_opt)
    stdscr.addstr(5,40, text_opt)

    
    stdscr.getch()

    # wait for user to click something
    while True: 
        if stdscr.getch() == curses.KEY_MOUSE:
            _, x, y, _, _ = curses.getmouse()

            if 0 <= x <= 15 and 3 <= y <= 10: 
                mode = '.pdf'

            elif 17 <= x <= 35 and 3 <= y <= 10:
                mode = '.docx'

            elif 37 <= x <= 55 and 3 <= y <= 10:
                mode = '.txt'

            else:
                continue

            return mode 
            
        else:
            continue

# lets user input file location
# displays error message if incorrect file type
def get_source(stdscr,mode):
    # clear out the terminal 
    stdscr.clear()

    # enable input mode and show the cursor
    curses.echo()  
    curses.curs_set(1)
    

    prompt = f"Step 2 of 5: Please enter the name of the file where the extract is stored (eg. article{mode}):"

    
    while True: 
        # display the prompt 
        stdscr.addstr(1, 1, prompt)  

        # get the user's input and decode it from a byte string to regular string
        source = stdscr.getstr().decode('utf-8')

        
        try: 
            validate_format(source, mode)
            break

        except ValueError:
            error_msg = 'Please ensure the file and its extension are entered correctly. Alternatively, press Ctrl-C to exit the program'

            show_error(stdscr, error_msg)
            
            # clear out the user's input, on the same row, and beginning column of input 
            stdscr.addstr(1, (1+len(prompt)), " " * (len(source)))

            stdscr.refresh() 
        
            continue
            

    # disable input mode and hide the cursor          
    curses.noecho()  
    curses.curs_set(0)

    # clear out the terminal 
    stdscr.clear()

    return source 


# lets user click to select output mode (.docx or .txt)
def get_dest(stdscr):

    # clear screen and hide the cursor
    stdscr.clear()
    curses.curs_set(0)

    # record all mouse movements and get the mouse position
    curses.mousemask(curses.ALL_MOUSE_EVENTS | curses.REPORT_MOUSE_POSITION)

    # option text to be displayed
    write_to_txt = "Write to text (.txt) file"
    write_to_docx = "Write to word (.docx) file"
    

    # display options on the screen
    stdscr.addstr(1,1, "Step 4 of 5: Please click select and click on an output method:")
    stdscr.addstr(3,1, write_to_txt)
    stdscr.addstr(3,35, write_to_docx)

    
    stdscr.getch()

    # wait for user to click something
    while True: 
        if stdscr.getch() == curses.KEY_MOUSE:
            _, x, y, _, _ = curses.getmouse()

            if 0 <= x <= 25 and 2 <= y <= 4: 
                dest = '.txt'

            elif 30 <= x <= 60 and 2 <= y <= 4:
                dest = '.docx'

            else:
                continue

            return dest
            
        else:
            continue

def get_title(stdscr):
    # clear out the terminal 
    stdscr.clear()

    # enable input mode and show the cursor
    curses.echo()  
    curses.curs_set(1)
    

    prompt = "Step 5 of 5: Please provide a title for the cloze passage:"

    # display the prompt 
    stdscr.addstr(1, 1, prompt)

    # get the user's input and decode it from a byte string to regular string
    title = stdscr.getstr().decode('utf-8')

    # disable input mode and hide the cursor          
    curses.noecho()  
    curses.curs_set(0)

    # clear out the terminal 
    stdscr.clear()

    return title
    
def get_limit(stdscr):
    # limit is blank by default
    limit = None

    # clear out the terminal 
    stdscr.clear()

    # enable input mode and show the cursor
    curses.echo()  
    curses.curs_set(1)
    
    stdscr.addstr(1,1, """Step 3 of 5: By default, the generator provides a blank for each sentence in the extract.
 If you would like to limit the number of blanks/sentences, please input a number (eg. 10). 
 Otherwise, press ENTER to continue.""")

    prompt = "Limit:"

    while True: 
        # display the prompt 
        stdscr.addstr(4, 1, prompt)

        # get the user's input and decode it from a byte string to regular string
        input = stdscr.getstr().decode('utf-8')

    
        if input == '':
            break
        
        else:
            try:
               validate_limit(input)
               limit = int(input)
               break
            
            except ValueError:
                error_msg = 'Please enter a valid positive number for the limit. Alternatively, press ENTER if you do not wish to set a limit.'

                show_error(stdscr, error_msg)
                
                # clear out the user's input, on the same row, and beginning column of input 
                stdscr.addstr(4, (1+len(prompt)), " " * (len(input)))

                stdscr.refresh()

                continue 
    
        
    # disable input mode and hide the cursor          
    curses.noecho()  
    curses.curs_set(0)

    # clear out the terminal 
    stdscr.clear()

    return limit

# shows error message(red) on stdscr 
def show_error(stdscr,error_msg):
    curses.start_color()

    # set red text against black background 
    curses.init_pair(1, curses.COLOR_RED, curses.COLOR_BLACK)

     # set the color pair to the error message text 
    stdscr.addstr(0,1, error_msg, curses.color_pair(1))

# shows success message (green) on stdscr
def show_success(stdscr, success_msg):
    curses.start_color()
     
    # set green text against black background 
    curses.init_pair(2, curses.COLOR_GREEN, curses.COLOR_BLACK)

    # set the color pair to the error message text 
    stdscr.addstr(0,1, success_msg, curses.color_pair(2))

# checks if limit is a positive integer
# else raises ValueError
def validate_limit(input):
    if int(input) <= 0:
        raise ValueError
    
# checks if extension of provided file matches selected mode
def validate_format(source, mode):
    
    ext = splitext(source)[1]

    if str(ext) != str(mode):
        raise ValueError

def read_txt(filename):
    with open(filename, 'r') as file:
        content = file.readlines()

        # join the list of paragraphs into a single string
        return '\n'.join(content)
    
def read_docx(filename):
    content = docx2txt.process(filename)
    return content 

def read_pdf(filename):
    content = extract_text(filename)

    # Removes unnecessary 'form feed' control character 
    # This is done as docx module dosen't accept control characters as input
    content = content.replace('\x0c', '')
    return content 



# matches chosen mode to method of content extraction
# runs that function (read_txt, read_docx or read_pdf)
# raises FileNotFoundError if file dosen't exist
def get_content(source, mode):

    accepted_modes = {
        '.txt': read_txt,
        '.docx': read_docx,
        '.pdf': read_pdf,
    }

    try:
        content = accepted_modes[mode](source)
        return content

    except FileNotFoundError:
        raise FileNotFoundError("File does not exist")

# global index to prefix blanks and total marks 
index = 0

# takes text as input and replaces text to form questions/answers      
def replace(content, limit):
    global index
    answer_sents = []
    question_sents = []

    nlp = spacy.load("en_core_web_sm")
            
    # call natural language processing object on str of text
    doc = nlp(content)

    for sent in list(doc.sents)[:limit]:
        # remove newlines from the end of sentence 
        text = sent.text.replace("\n", "")

        # randomise words in a sent for even distribution of blanks 
        randomized_sent = random.sample(list(sent), len(sent))

        for token in randomized_sent:
            if token.pos_ not in ['INTJ', 'NOUN', 'NUM', 'PUNCT', 'SYM', 'PROPN', 'X', 'SPACE']:

                # don't test students on contractions as they are too simplistic
                # \' uses an escape character to check if a literal apostrophe forms part of the token
                if '\'' not in token.orth_:
                    index += 1

                    answer = token
                    
                    blank = '_' * 20
                    replacement = f'{blank}[{answer.pos_}][{answer.text}]'

                    # find the answer in text and replace it with replacement 
                    # do this only for the first occurrence to avoid repetition
                    answer_sent = re.sub(r'\b' + answer.orth_ +  r'\b', f'({index}) {replacement}', text, count=1)
                    answer_sents.append(answer_sent)

                    question_sent = re.sub(r'\b' + answer.orth_ +  r'\b', f'({index}) {blank}', text, count=1)
                    question_sents.append(question_sent)
                    
                    break
                

    return question_sents, answer_sents

def generate_to_txt(content, title, limit):
    global index
    question_sents, answer_sents = replace(content, limit)
        
    with open('answers.txt', 'w') as file:
        file.write(f'[ANSWERS] Cloze Passage: {title} ({index} marks) \n')
        file.write('\n'.join(answer_sents))

    with open('questions.txt', 'w') as file:
        file.write(f'Cloze Passage: {title} ({index} marks) \n')
        file.write('\n'.join(question_sents))

def generate_to_docx(content, title, limit):
    global index 
    question_sents, answer_sents = replace(content, limit)

    # Create quesiton and answer Document objects
    questions = Document()
    answers = Document()

    # add titles to each of the documents
    # according to https://www.geeksforgeeks.org/working-with-titles-and-heading-python-docx-module/, 
    # level=1 is the largest section header. 
    questions.add_heading(f'Cloze Passage: {title} ({index} marks)\n', level=1)
    answers.add_heading(f'[ANSWERS] Cloze Passage: {title} ({index} marks) \n', level=1)
    

    # add a new paragraph to each of the documents 
    questions_paragraph = questions.add_paragraph((' '.join(question_sents)))
    answers_paragraph = answers.add_paragraph((' '.join(answer_sents)))

    # set double line spacing 
    questions_paragraph.paragraph_format.line_spacing = 2.0  
    answers_paragraph.paragraph_format.line_spacing = 2.0  

    # set justified text
    questions_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    answers_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # save the documents 
    questions.save('questions.docx')
    answers.save('answers.docx')

  
if __name__ == "__main__":
    # wrapper is necessary for any functions that use the stdscr from the curses module
    curses.wrapper(main)